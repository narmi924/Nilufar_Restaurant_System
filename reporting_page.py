#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
美莲花美食支出记录系统 - 报表页面
基于 Fluent Design 的功能页面，可嵌入到 NavigationInterface 中
"""

import sys
import pandas as pd
from PyQt6.QtCore import Qt, QDate, QTimer, QThread, pyqtSignal
from PyQt6.QtWidgets import (QApplication, QWidget, QVBoxLayout, QHBoxLayout, 
                             QHeaderView, QAbstractItemView, QTableWidgetItem, QFileDialog, QStackedWidget, QSplitter)
from qfluentwidgets import (TitleLabel, SubtitleLabel, BodyLabel, 
                            DateEdit, ComboBox, PrimaryPushButton, PushButton, 
                            TableWidget, CardWidget, InfoBar, InfoBarPosition, SegmentedWidget, TextEdit)
from PyQt6.QtGui import QColor

from models import get_expenses_by_date_range, get_all_categories, get_total_expenses_by_category
from workers import AIAnalysisWorker


class ReportingPage(QWidget):
    """
    报表页面类 - 使用 Fluent Design 风格
    提供单时段分析和双时段对比功能
    可嵌入到 NavigationInterface 中使用
    """
    
    # 定义自定义信号用于启动AI分析
    start_analysis_signal = pyqtSignal(list, tuple, list, tuple)
    
    def __init__(self, sound_manager=None, username=None, parent=None):
        """
        初始化报表页面
        
        Args:
            sound_manager: 声音管理器实例
            username: 当前登录用户名
            parent: 父控件
        """
        super().__init__(parent)
        self.sound_manager = sound_manager
        self.username = username
        self.categories = {}
        self.current_expenses_data = []
        # 用于存储双时段对比的数据，传递给AI分析
        self.period1_data = []
        self.period2_data = []
        self.p1_dates = None
        self.p2_dates = None
        self.init_ui()
        self.load_categories()
        self._connect_sound_effects()
        self.setup_ai_thread()
    
    def _connect_sound_effects(self):
        """
        连接下拉框音效
        """
        if self.sound_manager:
            try:
                # 为支出分类下拉框添加音效
                self.single_category.activated.connect(self.sound_manager.play_pop)
            except AttributeError:
                # 如果还没有创建下拉框，忽略
                pass
    
    def setup_ai_thread(self):
        """
        设置AI分析线程
        """
        # 创建线程和worker实例
        self.ai_thread = QThread()
        self.ai_worker = AIAnalysisWorker()
        
        # 将worker移动到新线程中
        self.ai_worker.moveToThread(self.ai_thread)
        
        # 连接信号与槽
        # worker的信号连接到UI更新方法
        self.ai_worker.finished.connect(self.on_ai_analysis_finished)
        self.ai_worker.error.connect(self.on_ai_analysis_error)
        
        # 自定义信号连接到worker的执行方法
        self.start_analysis_signal.connect(self.ai_worker.run_analysis)
        
        # 启动线程
        self.ai_thread.start()
    
    def on_ai_analysis_finished(self, report_text):
        """
        AI分析完成时的槽函数
        """
        self.analysis_text.setMarkdown(report_text)
        
        # 更新统计信息
        if hasattr(self, 'period1_data') and hasattr(self, 'period2_data'):
            self.update_comparison_summary_simple(
                self.period1_data, self.period2_data, 
                self.p1_dates[0], self.p1_dates[1], 
                self.p2_dates[0], self.p2_dates[1]
            )
        
        # 启用导出按钮
        self.export_btn.setEnabled(True)
        
        InfoBar.success(
            title="AI分析完成",
            content="DeepSeek智能分析报告已生成",
            duration=3000,
            position=InfoBarPosition.TOP,
            parent=self
        )
    
    def on_ai_analysis_error(self, error_message):
        """
        AI分析出错时的槽函数
        """
        self.analysis_text.setMarkdown(error_message)
        
        InfoBar.error(
            title="AI分析失败",
            content="请检查API配置和网络连接",
            duration=4000,
            position=InfoBarPosition.TOP,
            parent=self
        )
    
    def init_ui(self):
        """
        初始化用户界面 - 使用更优化的布局设计
        """
        # 创建主布局
        main_layout = QVBoxLayout(self)
        main_layout.setSpacing(10)
        main_layout.setContentsMargins(15, 10, 15, 10)
        
        # 创建上半部分：标题和查询条件（固定高度）
        self.create_header_section(main_layout)
        
        # 创建下半部分：报表展示区域（可伸缩）
        self.create_content_section(main_layout)
        
        # 初始加载主时段分析
        self.mode_segment.setCurrentItem("single")
        self.analysis_stack.setCurrentIndex(0)

    def create_header_section(self, parent_layout):
        """
        创建页面头部区域 - 标题和查询条件
        """
        header_card = CardWidget()
        header_layout = QVBoxLayout(header_card)
        header_layout.setContentsMargins(25, 20, 25, 20)
        header_layout.setSpacing(20)
        
        # 页面标题行
        title_layout = QHBoxLayout()
        main_title = SubtitleLabel("报表中心 /  دوكلات مەركىزى")
        main_title.setObjectName("pageTitle")
        title_layout.addWidget(main_title)
        title_layout.addStretch()
        header_layout.addLayout(title_layout)
        
        # 模式选择器
        mode_layout = QHBoxLayout()
        mode_label = BodyLabel("分析模式:")
        mode_label.setMinimumWidth(80)
        
        self.mode_segment = SegmentedWidget()
        self.mode_segment.addItem("single", "单时段分析")
        self.mode_segment.addItem("compare", "双时段对比")
        self.mode_segment.setFixedHeight(36)
        self.mode_segment.setMaximumWidth(240)
        self.mode_segment.currentItemChanged.connect(self.on_mode_changed)
        
        # 连接模式选择器音效
        if self.sound_manager:
            self.mode_segment.currentItemChanged.connect(lambda: self.sound_manager.play_pop())
        
        mode_layout.addWidget(mode_label)
        mode_layout.addWidget(self.mode_segment)
        mode_layout.addStretch()
        header_layout.addLayout(mode_layout)
        
        # 查询条件区域（动态切换）
        self.analysis_stack = QStackedWidget()
        self.create_single_analysis_widget()
        self.create_compare_analysis_widget()
        header_layout.addWidget(self.analysis_stack)
        
        parent_layout.addWidget(header_card)

    def create_single_analysis_widget(self):
        """
        创建单时段分析控件
        """
        widget = QWidget()
        layout = QVBoxLayout(widget)
        layout.setSpacing(15)
        layout.setContentsMargins(0, 0, 0, 0)
        
        # 查询条件行
        conditions_layout = QHBoxLayout()
        conditions_layout.setSpacing(25)
        
        # 开始日期
        start_label = BodyLabel("开始日期:")
        start_label.setMinimumWidth(70)
        self.single_start_date = DateEdit()
        self.single_start_date.setDate(QDate.currentDate().addDays(-30))
        self.single_start_date.setCalendarPopup(True)
        self.single_start_date.setMinimumWidth(200)
        self.single_start_date.setMaximumWidth(250)
        
        # 结束日期
        end_label = BodyLabel("结束日期:")
        end_label.setMinimumWidth(70)
        self.single_end_date = DateEdit()
        self.single_end_date.setDate(QDate.currentDate())
        self.single_end_date.setCalendarPopup(True)
        self.single_end_date.setMinimumWidth(200)
        self.single_end_date.setMaximumWidth(250)
        
        # 分类筛选
        category_label = BodyLabel("支出分类:")
        category_label.setMinimumWidth(70)
        self.single_category = ComboBox()
        self.single_category.addItem("全部分类", "0")
        self.single_category.setMinimumWidth(200)
        self.single_category.setMaximumWidth(280)
        
        # 生成按钮
        generate_btn = PrimaryPushButton("生成报表")
        generate_btn.setFixedSize(100, 32)
        generate_btn.clicked.connect(self.generate_single_report)
        
        # 单时段分析对所有用户开放，无需权限限制
        
        # 连接生成报表按钮音效
        if self.sound_manager:
            generate_btn.clicked.connect(self.sound_manager.play_click)
        
        conditions_layout.addWidget(start_label)
        conditions_layout.addWidget(self.single_start_date)
        conditions_layout.addWidget(end_label)  
        conditions_layout.addWidget(self.single_end_date)
        conditions_layout.addWidget(category_label)
        conditions_layout.addWidget(self.single_category)
        conditions_layout.addWidget(generate_btn)
        conditions_layout.addStretch()
        
        layout.addLayout(conditions_layout)
        self.analysis_stack.addWidget(widget)

    def create_compare_analysis_widget(self):
        """
        创建对比分析控件
        """
        widget = QWidget()
        layout = QVBoxLayout(widget)
        layout.setSpacing(15)
        layout.setContentsMargins(0, 0, 0, 0)
        
        # 时段1
        period1_layout = QHBoxLayout()
        period1_layout.setSpacing(25)
        
        p1_label = BodyLabel("时段1:")
        p1_label.setMinimumWidth(50)
        p1_start_label = BodyLabel("从")
        p1_start_label.setMinimumWidth(25)
        self.p1_start_date = DateEdit()
        self.p1_start_date.setDate(QDate.currentDate().addDays(-30))
        self.p1_start_date.setCalendarPopup(True)
        self.p1_start_date.setMinimumWidth(200)
        self.p1_start_date.setMaximumWidth(250)
        
        p1_end_label = BodyLabel("到")
        p1_end_label.setMinimumWidth(25)
        self.p1_end_date = DateEdit()
        self.p1_end_date.setDate(QDate.currentDate().addDays(-15))
        self.p1_end_date.setCalendarPopup(True)
        self.p1_end_date.setMinimumWidth(200)
        self.p1_end_date.setMaximumWidth(250)
        
        period1_layout.addWidget(p1_label)
        period1_layout.addWidget(p1_start_label)
        period1_layout.addWidget(self.p1_start_date)
        period1_layout.addWidget(p1_end_label)
        period1_layout.addWidget(self.p1_end_date)
        period1_layout.addStretch()
        
        # 时段2
        period2_layout = QHBoxLayout()
        period2_layout.setSpacing(25)
        
        p2_label = BodyLabel("时段2:")
        p2_label.setMinimumWidth(50)
        p2_start_label = BodyLabel("从")
        p2_start_label.setMinimumWidth(25)
        self.p2_start_date = DateEdit()
        self.p2_start_date.setDate(QDate.currentDate().addDays(-14))
        self.p2_start_date.setCalendarPopup(True)
        self.p2_start_date.setMinimumWidth(200)
        self.p2_start_date.setMaximumWidth(250)
        
        p2_end_label = BodyLabel("到")
        p2_end_label.setMinimumWidth(25)
        self.p2_end_date = DateEdit()
        self.p2_end_date.setDate(QDate.currentDate())
        self.p2_end_date.setCalendarPopup(True)
        self.p2_end_date.setMinimumWidth(200)
        self.p2_end_date.setMaximumWidth(250)
        
        # 生成按钮
        compare_btn = PrimaryPushButton("生成对比")
        compare_btn.setFixedSize(100, 32)
        compare_btn.clicked.connect(self.generate_compare_report)
        
        # 如果不是管理员，显示权限提示
        if self.username != "admin":
            compare_btn.setToolTip("AI分析功能仅限管理员使用\n请使用admin账户登录")
        
        # 连接生成对比按钮音效
        if self.sound_manager:
            compare_btn.clicked.connect(self.sound_manager.play_click)
        
        period2_layout.addWidget(p2_label)
        period2_layout.addWidget(p2_start_label)
        period2_layout.addWidget(self.p2_start_date)
        period2_layout.addWidget(p2_end_label)
        period2_layout.addWidget(self.p2_end_date)
        period2_layout.addWidget(compare_btn)
        period2_layout.addStretch()
        
        layout.addLayout(period1_layout)
        layout.addLayout(period2_layout)
        self.analysis_stack.addWidget(widget)

    def create_content_section(self, parent_layout):
        """
        创建内容区域 - 统计信息和报表表格
        """
        # 使用水平分割器实现灵活布局
        splitter = QSplitter(Qt.Orientation.Horizontal)
        
        # 左侧：统计信息和操作按钮
        self.create_stats_panel(splitter)
        
        # 右侧：报表表格
        self.create_table_panel(splitter)
        
        # 设置分割比例 (25% : 75%)
        splitter.setSizes([300, 900])
        splitter.setChildrenCollapsible(False)
        
        parent_layout.addWidget(splitter, 1)  # 设置伸缩因子为1

    def create_stats_panel(self, parent_splitter):
        """
        创建统计信息面板
        """
        stats_card = CardWidget()
        stats_layout = QVBoxLayout(stats_card)
        stats_layout.setContentsMargins(20, 20, 20, 20)
        stats_layout.setSpacing(20)
        
        # 统计信息标题
        stats_title = SubtitleLabel("统计摘要")
        stats_layout.addWidget(stats_title)
        
        # 统计数据显示区域
        self.stats_container = QWidget()
        stats_data_layout = QVBoxLayout(self.stats_container)
        stats_data_layout.setSpacing(15)
        
        # 基本统计信息
        self.total_records_label = BodyLabel("记录数: --")
        self.total_amount_label = BodyLabel("总金额: --")
        self.avg_amount_label = BodyLabel("平均金额: --")
        self.date_range_label = BodyLabel("时间范围: --")
        
        stats_data_layout.addWidget(self.total_records_label)
        stats_data_layout.addWidget(self.total_amount_label)
        stats_data_layout.addWidget(self.avg_amount_label)
        stats_data_layout.addWidget(self.date_range_label)
        
        stats_layout.addWidget(self.stats_container)
        
        # 分隔线
        stats_layout.addWidget(BodyLabel(""))
        
        # 操作按钮区域
        self.create_action_buttons(stats_layout)
        
        stats_layout.addStretch()
        parent_splitter.addWidget(stats_card)

    def create_action_buttons(self, parent_layout):
        """
        创建操作按钮
        """
        buttons_title = BodyLabel("操作:")
        buttons_title.setStyleSheet("font-weight: bold; color: black;")
        parent_layout.addWidget(buttons_title)
        
        # 生成饼图按钮
        self.pie_chart_btn = PushButton("生成饼图")
        self.pie_chart_btn.setFixedHeight(36)
        self.pie_chart_btn.setStyleSheet("background-color: #1A5319; color: white; border: none; border-radius: 6px;")
        self.pie_chart_btn.clicked.connect(self.generate_pie_chart)
        self.pie_chart_btn.setEnabled(False)
        
        # 连接生成饼图按钮音效
        if self.sound_manager:
            self.pie_chart_btn.clicked.connect(self.sound_manager.play_click)
        
        # 导出报表按钮
        self.export_btn = PushButton("导出单时段报表")
        self.export_btn.setFixedHeight(36)
        self.export_btn.setStyleSheet("background-color: #1A5319; color: white; border: none; border-radius: 6px;")
        self.export_btn.clicked.connect(self.export_report)
        self.export_btn.setEnabled(False)
        
        # 连接导出报表按钮音效
        if self.sound_manager:
            self.export_btn.clicked.connect(self.sound_manager.play_click)
        
        # 重置条件按钮
        self.reset_btn = PushButton("重置条件")
        self.reset_btn.setFixedHeight(36)
        self.reset_btn.setStyleSheet("background-color: #1A5319; color: white; border: none; border-radius: 6px;")
        self.reset_btn.clicked.connect(self.reset_conditions)
        
        # 连接重置条件按钮音效
        if self.sound_manager:
            self.reset_btn.clicked.connect(self.sound_manager.play_click)
        
        parent_layout.addWidget(self.pie_chart_btn)
        parent_layout.addWidget(self.export_btn)
        parent_layout.addWidget(self.reset_btn)

    def create_table_panel(self, parent_splitter):
        """
        创建表格面板（支持表格和文本显示切换）
        """
        table_card = CardWidget()
        table_layout = QVBoxLayout(table_card)
        table_layout.setContentsMargins(20, 20, 20, 20)
        table_layout.setSpacing(15)
        
        # 表格标题
        table_header = QHBoxLayout()
        self.table_title = SubtitleLabel("报表详情")
        table_header.addWidget(self.table_title)
        table_header.addStretch()
        table_layout.addLayout(table_header)
        
        # 创建堆叠式容器用于切换显示
        self.content_stack = QStackedWidget()
        
        # 创建表格容器
        table_container = QWidget()
        table_container_layout = QVBoxLayout(table_container)
        table_container_layout.setContentsMargins(0, 0, 0, 0)
        
        self.report_table = TableWidget()
        self.report_table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self.report_table.setAlternatingRowColors(True)
        self.report_table.setSortingEnabled(True)
        self.report_table.setMinimumHeight(400)  # 减少高度为总计行留空间
        table_container_layout.addWidget(self.report_table)
        
        # 创建固定的总计行
        self.create_fixed_total_row_report(table_container_layout)
        
        # 创建文本分析容器
        text_container = QWidget()
        text_container_layout = QVBoxLayout(text_container)
        text_container_layout.setContentsMargins(0, 0, 0, 0)
        
        self.analysis_text = TextEdit()
        self.analysis_text.setReadOnly(True)
        
        # 根据用户权限设置初始提示文本
        if self.username == "admin":
            initial_text = "请点击'生成对比'按钮获取AI智能分析报告..."
        else:
            initial_text = "🔒 AI智能分析功能仅限管理员使用\n\n如需使用此功能，请使用admin账户登录。\n\n💡 您仍可以使用'单时段分析'功能查看基础报表数据。"
        
        self.analysis_text.setPlainText(initial_text)
        text_container_layout.addWidget(self.analysis_text)
        
        # 添加到堆叠容器
        self.content_stack.addWidget(table_container)  # 索引0：表格
        self.content_stack.addWidget(text_container)   # 索引1：文本分析
        
        table_layout.addWidget(self.content_stack)
        parent_splitter.addWidget(table_card)
    
    def create_fixed_total_row_report(self, parent_layout):
        """
        创建固定在底部的总计行（报表页面）
        """
        # 创建总计行容器
        total_container = QWidget()
        total_container.setFixedHeight(40)
        total_container.setStyleSheet("background-color: lightgray; border: 1px solid #ccc;")
        
        total_layout = QHBoxLayout(total_container)
        total_layout.setContentsMargins(0, 0, 0, 0)
        total_layout.setSpacing(0)
        
        # 创建与表格列对应的标签，使用固定宽度与表格列对齐
        # 注意：报表页面的列数和宽度可能会动态变化，所以使用通用的布局
        self.report_total_col1_label = BodyLabel("总计")
        self.report_total_col1_label.setMinimumWidth(150)
        self.report_total_col1_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.report_total_col1_label.setStyleSheet("border-right: 1px solid #ccc; padding: 5px;")
        
        self.report_total_col2_label = BodyLabel("جەمئىي - 0条记录")
        self.report_total_col2_label.setMinimumWidth(200)
        self.report_total_col2_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.report_total_col2_label.setStyleSheet("border-right: 1px solid #ccc; padding: 5px;")
        
        self.report_total_col3_label = BodyLabel("¥0.00")
        self.report_total_col3_label.setMinimumWidth(100)
        self.report_total_col3_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.report_total_col3_label.setStyleSheet("border-right: 1px solid #ccc; padding: 5px; font-weight: bold;")
        
        self.report_total_col4_label = BodyLabel("¥0.00")
        self.report_total_col4_label.setMinimumWidth(150)
        self.report_total_col4_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.report_total_col4_label.setStyleSheet("border-right: 1px solid #ccc; padding: 5px;")
        
        self.report_total_col5_label = BodyLabel("0 خاتىرە")
        self.report_total_col5_label.setMinimumWidth(120)
        self.report_total_col5_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.report_total_col5_label.setStyleSheet("padding: 5px;")
        
        total_layout.addWidget(self.report_total_col1_label)
        total_layout.addWidget(self.report_total_col2_label)
        total_layout.addWidget(self.report_total_col3_label)
        total_layout.addWidget(self.report_total_col4_label)
        total_layout.addWidget(self.report_total_col5_label)
        total_layout.addStretch()  # 添加弹性空间以适应不同的列数
        
        parent_layout.addWidget(total_container)

    def on_mode_changed(self, key):
        """
        处理模式切换
        """
        index = 0 if key == "single" else 1
        self.analysis_stack.setCurrentIndex(index)
        
        # 切换内容显示区域
        if hasattr(self, 'content_stack'):
            if key == "single":
                self.content_stack.setCurrentIndex(0)  # 显示表格
                self.table_title.setText("报表详情")
            else:
                self.content_stack.setCurrentIndex(1)  # 显示文本分析
                self.table_title.setText("对比分析报告")
        
        # 控制饼图按钮的显示（仅在单时段分析模式下显示）
        if hasattr(self, 'pie_chart_btn'):
            self.pie_chart_btn.setVisible(key == "single")
        
        # 更新导出按钮文字
        if hasattr(self, 'export_btn'):
            if key == "single":
                self.export_btn.setText("导出单时段报表")
            else:
                self.export_btn.setText("导出双时段对比报告")

    def load_categories(self):
        """
        加载分类数据
        """
        try:
            categories = get_all_categories()
            self.categories = {cat['id']: cat for cat in categories}
            
            # 更新分类下拉框
            self.single_category.clear()
            self.single_category.addItem("全部分类", "0")
            
            for category in categories:
                display_text = f"{category['name_cn']} / {category['name_ug']}"
                self.single_category.addItem(display_text, str(category['id']))
                
        except Exception as e:
            InfoBar.error(
                title="加载失败",
                content=f"无法加载分类数据：{str(e)}",
                duration=3000,
                position=InfoBarPosition.TOP,
                parent=self
            )

    def reset_conditions(self):
        """
        重置条件 - 将时间段重置为默认值
        """
        current_mode = self.mode_segment.currentItem()
        
        if current_mode == "single":
            # 重置单时段分析的条件
            self.single_start_date.setDate(QDate.currentDate().addDays(-30))
            self.single_end_date.setDate(QDate.currentDate())
            self.single_category.setCurrentIndex(0)  # 重置为"全部分类"
            
            # 清空表格和统计信息
            self.report_table.setRowCount(0)
            self.total_records_label.setText("记录数: --")
            self.total_amount_label.setText("总金额: --")
            self.avg_amount_label.setText("平均金额: --")
            self.date_range_label.setText("时间范围: --")
            
            # 重置固定总计行
            self.report_total_col2_label.setText("جەمئىي - 0条记录")
            self.report_total_col3_label.setText("¥0.00")
            self.report_total_col4_label.setText("¥0.00")
            self.report_total_col5_label.setText("0 خاتىرە")
            
            # 禁用按钮
            self.pie_chart_btn.setEnabled(False)
            self.export_btn.setEnabled(False)
            
        else:
            # 重置对比分析的条件
            self.p1_start_date.setDate(QDate.currentDate().addDays(-30))
            self.p1_end_date.setDate(QDate.currentDate().addDays(-15))
            self.p2_start_date.setDate(QDate.currentDate().addDays(-14))
            self.p2_end_date.setDate(QDate.currentDate())
            
            # 清空分析文本和统计信息
            self.analysis_text.setPlainText("请点击'生成对比'按钮获取分析报告...")
            self.total_records_label.setText("记录数: --")
            self.total_amount_label.setText("总金额: --")
            self.avg_amount_label.setText("平均金额: --")
            self.date_range_label.setText("时间范围: --")
            
            # 禁用按钮
            self.pie_chart_btn.setEnabled(False)
            self.export_btn.setEnabled(False)
        
        # 清空当前数据
        if hasattr(self, 'current_expenses_data'):
            self.current_expenses_data = []
        
        InfoBar.success(
            title="重置完成",
            content="查询条件已重置为默认值",
            duration=2000,
            position=InfoBarPosition.TOP,
            parent=self
        )

    def generate_single_report(self):
        """
        生成单时段分析报表
        """
        try:
            # 获取查询条件
            start_date = self.single_start_date.date().toString("yyyy-MM-dd")
            end_date = self.single_end_date.date().toString("yyyy-MM-dd")
            
            # 验证日期范围
            if self.single_start_date.date() > self.single_end_date.date():
                InfoBar.warning(
                    title="日期错误",
                    content="开始日期不能晚于结束日期",
                    duration=3000,
                    position=InfoBarPosition.TOP,
                    parent=self
                )
                return
            
            # 获取分类筛选 - 使用备用方案处理QFluentWidgets ComboBox的data问题
            current_index = self.single_category.currentIndex()
            category_id_str = self.single_category.currentData()
            
            if category_id_str is None:
                category_id_str = "0"
            
            # 获取当前选择的分类信息
            selected_category_text = self.single_category.currentText()
            
            # 备用方案：如果选择了非全部分类但ID仍为0，通过文本匹配获取正确的分类ID
            if current_index > 0 and category_id_str == "0":
                categories = get_all_categories()
                for category in categories:
                    # 精确匹配：检查分类名称是否与选择的文本完全匹配
                    display_text = f"{category['name_cn']} / {category['name_ug']}"
                    if display_text == selected_category_text:
                        category_id_str = str(category['id'])
                        break
            
            # 查询数据
            if category_id_str == "0":
                raw_expenses = get_expenses_by_date_range(start_date, end_date)
                report_type = "全部分类"
            else:
                category_id = int(category_id_str)
                raw_expenses = get_expenses_by_date_range(start_date, end_date, category_id)
                report_type = f"分类筛选: {selected_category_text}"
            
            # 转换元组数据为字典格式
            expenses = []
            for row in raw_expenses:
                expenses.append({
                    'date': row[0],           # expense_date
                    'category_name': row[1],  # category name_cn
                    'amount': row[2],         # amount
                    'operator': row[3],       # username
                    'description': row[4] if row[4] else ''  # notes
                })
            
            # 更新统计信息和表格
            self.update_summary(expenses, start_date, end_date)
            self.update_main_table(expenses)
            
            # 保存当前报表数据用于饼图生成
            self.current_expenses_data = expenses
            
            # 启用按钮
            self.pie_chart_btn.setEnabled(len(expenses) > 0)
            self.export_btn.setEnabled(len(expenses) > 0)
            
            InfoBar.success(
                title="生成完成",
                content=f"报表生成成功 ({report_type})，共 {len(expenses)} 条记录",
                duration=2000,
                position=InfoBarPosition.TOP,
                parent=self
            )
            
        except Exception as e:
            InfoBar.error(
                title="生成失败",
                content=f"报表生成失败：{str(e)}",
                duration=4000,
                position=InfoBarPosition.TOP,
                parent=self
            )

    def generate_compare_report(self):
        """
        生成双时段对比报表
        """
        try:
            # 权限检查：只有管理员才能使用AI分析功能
            if self.username != "admin":
                self.analysis_text.setPlainText("❌ 权限不足\n\n🔒 AI智能分析功能仅限管理员使用\n\n💡 原因说明：\n• AI分析会消耗API tokens，需要严格控制使用权限\n• 确保系统资源的合理分配和成本控制\n\n📞 如需使用此功能，请联系系统管理员或使用管理员账户登录。")
                
                InfoBar.warning(
                    title="权限不足",
                    content="请使用管理员账户登录以使用AI分析功能",
                    duration=4000,
                    position=InfoBarPosition.TOP,
                    parent=self
                )
                return
            
            # 获取两个时段的日期
            p1_start = self.p1_start_date.date().toString("yyyy-MM-dd")
            p1_end = self.p1_end_date.date().toString("yyyy-MM-dd")
            p2_start = self.p2_start_date.date().toString("yyyy-MM-dd")
            p2_end = self.p2_end_date.date().toString("yyyy-MM-dd")
            
            # 验证日期范围
            if (self.p1_start_date.date() > self.p1_end_date.date() or 
                self.p2_start_date.date() > self.p2_end_date.date()):
                InfoBar.warning(
                    title="日期错误",
                    content="开始日期不能晚于结束日期",
                    duration=3000,
                    position=InfoBarPosition.TOP,
                    parent=self
                )
                return
            
            # 查询两个时段的数据并按分类汇总
            from models import get_total_expenses_by_category
            
            # 获取时段1的分类汇总数据
            try:
                p1_category_data = get_total_expenses_by_category(p1_start, p1_end)
                if p1_category_data is None:
                    p1_category_data = []
            except Exception as e:
                InfoBar.error(
                    title="数据查询失败",
                    content=f"查询时段1数据失败：{str(e)}",
                    duration=4000,
                    position=InfoBarPosition.TOP,
                    parent=self
                )
                return
            
            # 处理时段1数据 - 注意get_total_expenses_by_category返回字典格式
            self.period1_data = []
            for item in p1_category_data:
                if item and isinstance(item, dict):
                    self.period1_data.append({
                        'category_cn': str(item.get('category_name', '未知分类')),
                        'total_amount': float(item.get('total_amount', 0.0))
                    })
            
            # 获取时段2的分类汇总数据
            try:
                p2_category_data = get_total_expenses_by_category(p2_start, p2_end)
                if p2_category_data is None:
                    p2_category_data = []
            except Exception as e:
                InfoBar.error(
                    title="数据查询失败",
                    content=f"查询时段2数据失败：{str(e)}",
                    duration=4000,
                    position=InfoBarPosition.TOP,
                    parent=self
                )
                return
            
            # 处理时段2数据 - 注意get_total_expenses_by_category返回字典格式
            self.period2_data = []
            for item in p2_category_data:
                if item and isinstance(item, dict):
                    self.period2_data.append({
                        'category_cn': str(item.get('category_name', '未知分类')),
                        'total_amount': float(item.get('total_amount', 0.0))
                    })
            
            # 保存日期信息到实例变量
            self.p1_dates = (p1_start, p1_end)
            self.p2_dates = (p2_start, p2_end)
            
            # 检查是否有数据
            if not self.period1_data and not self.period2_data:
                InfoBar.warning(
                    title="无数据",
                    content="选择的时间段内没有支出记录，无法生成对比分析",
                    duration=3000,
                    position=InfoBarPosition.TOP,
                    parent=self
                )
                self.analysis_text.setPlainText("选择的时间段内没有支出记录，无法生成对比分析。\n\n请选择有数据的时间段进行对比。")
                return
            
            # 验证数据有效性
            if not self.period1_data and not self.period2_data:
                self.analysis_text.setPlainText("❌ 数据验证失败：两个时段都没有支出记录，无法进行对比分析。")
                InfoBar.warning(
                    title="数据不足",
                    content="两个时段都没有支出记录",
                    duration=3000,
                    position=InfoBarPosition.TOP,
                    parent=self
                )
                return
            elif not self.period1_data:
                self.analysis_text.setPlainText("❌ 数据验证失败：时段1没有支出记录。")
                InfoBar.warning(
                    title="数据不足",
                    content="时段1没有支出记录",
                    duration=3000,
                    position=InfoBarPosition.TOP,
                    parent=self
                )
                return
            elif not self.period2_data:
                self.analysis_text.setPlainText("❌ 数据验证失败：时段2没有支出记录。")
                InfoBar.warning(
                    title="数据不足", 
                    content="时段2没有支出记录",
                    duration=3000,
                    position=InfoBarPosition.TOP,
                    parent=self
                )
                return
            
            # 显示正在分析的提示
            self.analysis_text.setPlainText("🤖 正在请求DeepSeek AI进行深度分析...\n\n📊 分析内容：\n• 总体支出对比\n• 核心品类异动分析\n• 潜在风险与机遇预警\n• 经营诊断与具体建议\n\n⏱️ 预计耗时：30-90秒（复杂分析需要更多时间）\n\n✨ 界面保持响应，请耐心等待...\n\n💡 提示：如果超时，您的API调用仍然有效，请稍后重试")
            
            # 使用多线程方式启动AI分析
            self.start_analysis_signal.emit(
                self.period1_data, self.p1_dates, 
                self.period2_data, self.p2_dates
            )
            
            # 禁用饼图按钮（对比模式不支持饼图）
            self.pie_chart_btn.setEnabled(False)
            
        except Exception as e:
            InfoBar.error(
                title="对比失败",
                content=f"对比分析失败：{str(e)}",
                duration=4000,
                position=InfoBarPosition.TOP,
                parent=self
            )
    
    def closeEvent(self, event):
        """
        窗口关闭事件 - 确保线程正确退出
        """
        if hasattr(self, 'ai_thread'):
            self.ai_thread.quit()
            self.ai_thread.wait()
        super().closeEvent(event)
    


    def update_summary(self, expenses, start_date, end_date):
        """
        更新统计摘要信息
        """
        total_records = len(expenses)
        total_amount = sum(expense['amount'] for expense in expenses)
        avg_amount = total_amount / total_records if total_records > 0 else 0
        
        # 更新标签
        self.total_records_label.setText(f"记录数: {total_records}")
        self.total_amount_label.setText(f"总金额: ¥{total_amount:.2f}")
        self.avg_amount_label.setText(f"平均金额: ¥{avg_amount:.2f}")
        self.date_range_label.setText(f"时间范围: {start_date} 至 {end_date}")

    def update_comparison_summary(self, p1_expenses, p2_expenses, p1_start, p1_end, p2_start, p2_end):
        """
        更新对比分析的统计摘要
        """
        p1_total = sum(expense['amount'] for expense in p1_expenses)
        p2_total = sum(expense['amount'] for expense in p2_expenses)
        change_amount = p2_total - p1_total
        change_pct = (change_amount / p1_total * 100) if p1_total > 0 else 0
        
        self.total_records_label.setText(f"时段1: {len(p1_expenses)} 条 | 时段2: {len(p2_expenses)} 条")
        self.total_amount_label.setText(f"时段1: ¥{p1_total:.2f} | 时段2: ¥{p2_total:.2f}")
        self.avg_amount_label.setText(f"变化: ¥{change_amount:.2f} ({change_pct:+.1f}%)")
        self.date_range_label.setText(f"对比时段: {p1_start}~{p1_end} vs {p2_start}~{p2_end}")

    def update_comparison_summary_simple(self, period1_data, period2_data, p1_start, p1_end, p2_start, p2_end):
        """
        更新对比分析的简化统计摘要
        """
        p1_total = sum(item['total_amount'] for item in period1_data)
        p2_total = sum(item['total_amount'] for item in period2_data)
        change_amount = p2_total - p1_total
        change_pct = (change_amount / p1_total * 100) if p1_total > 0 else 0
        
        p1_categories = len(period1_data)
        p2_categories = len(period2_data)
        
        self.total_records_label.setText(f"时段1分类: {p1_categories} 个 | 时段2分类: {p2_categories} 个")
        self.total_amount_label.setText(f"时段1总额: ¥{p1_total:.2f} | 时段2总额: ¥{p2_total:.2f}")
        self.avg_amount_label.setText(f"变化金额: ¥{change_amount:.2f} ({change_pct:+.1f}%)")
        self.date_range_label.setText(f"对比时段: {p1_start}~{p1_end} vs {p2_start}~{p2_end}")

    def update_main_table(self, expenses):
        """
        更新主报表表格 - 使用与每日管理页面一致的简雅风格，分类列显示双语，删除ID列
        """
        # 设置表格列
        headers = ["日期", "分类", "金额", "备注", "操作员"]  # 删除ID列
        self.report_table.setColumnCount(len(headers))
        self.report_table.setHorizontalHeaderLabels(headers)
        
        # 设置行数：只显示数据行，不包含总计行
        self.report_table.setRowCount(len(expenses))
        
        # 设置表格属性 - 与每日管理页面保持一致
        self.report_table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self.report_table.setAlternatingRowColors(True)
        self.report_table.setSortingEnabled(True)
        self.report_table.setMinimumHeight(400)  # 减少高度为总计行留空间
        
        # 填充数据行
        for row, expense in enumerate(expenses):
            # 获取维语分类名称
            category_ug = self.get_category_ug_name(expense['category_name'])
            category_display = f"{expense['category_name']} / {category_ug}"
            
            self.report_table.setItem(row, 0, QTableWidgetItem(str(expense['date'])))
            self.report_table.setItem(row, 1, QTableWidgetItem(category_display))
            self.report_table.setItem(row, 2, QTableWidgetItem(f"¥{expense['amount']:.2f}"))
            self.report_table.setItem(row, 3, QTableWidgetItem(expense.get('description', '')))
            self.report_table.setItem(row, 4, QTableWidgetItem(expense.get('operator', '')))  # 删除ID列
        
        # 更新固定的总计行显示
        self.update_fixed_total_row_report(expenses)
        
        # 设置列宽 - 与每日管理页面一致
        header = self.report_table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.ResizeMode.Fixed)  # 日期
        header.setSectionResizeMode(1, QHeaderView.ResizeMode.Fixed)  # 分类
        header.setSectionResizeMode(2, QHeaderView.ResizeMode.Fixed)  # 金额
        header.setSectionResizeMode(3, QHeaderView.ResizeMode.Fixed)  # 备注
        header.setSectionResizeMode(4, QHeaderView.ResizeMode.Fixed)  # 操作员
        
        self.report_table.setColumnWidth(0, 150)  # 日期
        self.report_table.setColumnWidth(1, 200)  # 分类
        self.report_table.setColumnWidth(2, 100)  # 金额
        self.report_table.setColumnWidth(3, 200)  # 备注
        self.report_table.setColumnWidth(4, 150)  # 操作员
    
    def update_fixed_total_row_report(self, expenses):
        """
        更新固定的总计行显示（报表页面）
        """
        total_records = len(expenses)
        total_amount = sum(expense['amount'] for expense in expenses) if expenses else 0.0
        
        # 更新各个标签的文本
        self.report_total_col2_label.setText(f"جەمئىي - {total_records}条记录")
        self.report_total_col3_label.setText(f"¥{total_amount:.2f}")
        self.report_total_col4_label.setText(f"¥{total_amount:.2f}")
        self.report_total_col5_label.setText(f"{total_records} خاتىرە")

    def update_dual_table(self, comparison_data):
        """
        更新双时段对比表格 - 分类列显示双语
        """
        # 设置表格列
        headers = ["分类", "时段1金额", "时段1次数", "时段2金额", "时段2次数", "金额变化", "变化%"]
        self.report_table.setColumnCount(len(headers))
        self.report_table.setHorizontalHeaderLabels(headers)
        
        # 设置行数
        self.report_table.setRowCount(len(comparison_data))
        
        # 填充数据
        for row, data in enumerate(comparison_data):
            # 获取维语分类名称
            category_ug = self.get_category_ug_name(data['category'])
            category_display = f"{data['category']} / {category_ug}"
            
            self.report_table.setItem(row, 0, QTableWidgetItem(category_display))
            self.report_table.setItem(row, 1, QTableWidgetItem(f"¥{data['p1_amount']:.2f}"))
            self.report_table.setItem(row, 2, QTableWidgetItem(str(data['p1_count'])))
            self.report_table.setItem(row, 3, QTableWidgetItem(f"¥{data['p2_amount']:.2f}"))
            self.report_table.setItem(row, 4, QTableWidgetItem(str(data['p2_count'])))
            
            # 金额变化（带颜色指示）
            change_item = QTableWidgetItem(f"¥{data['amount_change']:+.2f}")
            if data['amount_change'] > 0:
                change_item.setForeground(QColor(220, 53, 69))  # 增加用红色
            elif data['amount_change'] < 0:
                change_item.setForeground(QColor(40, 167, 69))  # 减少用绿色
            self.report_table.setItem(row, 5, change_item)
            
            # 变化百分比
            pct_item = QTableWidgetItem(f"{data['amount_change_pct']:+.1f}%")
            if data['amount_change_pct'] > 0:
                pct_item.setForeground(QColor(220, 53, 69))
            elif data['amount_change_pct'] < 0:
                pct_item.setForeground(QColor(40, 167, 69))
            self.report_table.setItem(row, 6, pct_item)
        
        # 调整列宽
        header = self.report_table.horizontalHeader()
        for i in range(len(headers)):
            header.setSectionResizeMode(i, QHeaderView.ResizeMode.ResizeToContents)

    def export_report(self):
        """
        导出报表文件
        """
        try:
            # 使用堆叠控件的当前索引来判断模式，更可靠
            current_index = self.analysis_stack.currentIndex()
            
            if current_index == 0:
                # 单时段分析 - 导出Excel
                self.export_single_report()
            else:
                # 双时段对比 - 导出Word文档
                self.export_compare_report()
                
        except Exception as e:
            InfoBar.error(
                title="导出失败",
                content=f"导出时发生错误：{str(e)}",
                duration=4000,
                position=InfoBarPosition.TOP,
                parent=self
            )
    
    def export_single_report(self):
        """
        导出单时段分析报表到Excel
        """
        # 检查是否有数据
        if self.report_table.rowCount() == 0:
            InfoBar.warning(
                title="导出提醒",
                content="当前没有数据可以导出",
                duration=3000,
                position=InfoBarPosition.TOP,
                parent=self
            )
            return
        
        # 获取时间段
        start_date = self.single_start_date.date().toString("yyyy-MM-dd")
        end_date = self.single_end_date.date().toString("yyyy-MM-dd")
        filename = f"美莲花美食_{start_date}_至_{end_date}_单时段报表.xlsx"
        
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "导出单时段报表",
            filename,
            "Excel文件 (*.xlsx)"
        )
        
        if not file_path:
            return
        
        # 创建DataFrame
        headers = []
        for col in range(self.report_table.columnCount()):
            headers.append(self.report_table.horizontalHeaderItem(col).text())
        
        data = []
        for row in range(self.report_table.rowCount()):
            row_data = []
            for col in range(self.report_table.columnCount()):
                item = self.report_table.item(row, col)
                row_data.append(item.text() if item else "")
            data.append(row_data)
        
        # 添加总计行
        if hasattr(self, 'current_expenses_data') and self.current_expenses_data:
            total_records = len(self.current_expenses_data)
            total_amount = sum(expense['amount'] for expense in self.current_expenses_data)
            data.append([
                "总计",
                f"جەمئىي - {total_records}条记录",
                f"¥{total_amount:.2f}",
                f"¥{total_amount:.2f}",
                f"{total_records} خاتىرە"
            ])
        
        df = pd.DataFrame(data, columns=headers)
        
        # 导出到Excel
        with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='报表数据', index=False)
            
            # 设置列宽
            worksheet = writer.sheets['报表数据']
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 30)
                worksheet.column_dimensions[column_letter].width = adjusted_width
        
        InfoBar.success(
            title="导出成功",
            content=f"单时段报表已成功导出到：\n{file_path}",
            duration=3000,
            position=InfoBarPosition.TOP,
            parent=self
        )
    
    def export_compare_report(self):
        """
        导出双时段对比报告到Word文档
        """
        # 检查是否有分析内容
        analysis_content = self.analysis_text.toPlainText()
        if not analysis_content or analysis_content == "请点击'生成对比'按钮获取分析报告...":
            InfoBar.warning(
                title="导出提醒",
                content="请先生成双时段对比分析报告",
                duration=3000,
                position=InfoBarPosition.TOP,
                parent=self
            )
            return
        
        # 获取时间段
        p1_start = self.p1_start_date.date().toString("yyyy-MM-dd")
        p1_end = self.p1_end_date.date().toString("yyyy-MM-dd")
        p2_start = self.p2_start_date.date().toString("yyyy-MM-dd")
        p2_end = self.p2_end_date.date().toString("yyyy-MM-dd")
        filename = f"美莲花美食_{p1_start}_至_{p1_end}_对比_{p2_start}_至_{p2_end}_双时段对比报告.docx"
        
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "导出双时段对比报告",
            filename,
            "Word文档 (*.docx)"
        )
        
        if not file_path:
            return
        
        try:
            # 尝试导入python-docx
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.oxml.ns import qn
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # 创建Word文档
            doc = Document()
            
            # 设置文档的默认字体为支持中文和维语的字体
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial Unicode MS'  # 支持多语言的字体
            font.size = Pt(12)
            
            # 设置多语言字体支持
            style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')  # 中文字体
            style.element.rPr.rFonts.set(qn('w:ascii'), 'Arial Unicode MS')  # 英文和其他字符
            style.element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial Unicode MS')  # 高位ANSI字符
            style.element.rPr.rFonts.set(qn('w:cs'), 'Arial Unicode MS')  # 复杂脚本（包括阿拉伯文、维语等）
            
            # 添加标题
            title_para = doc.add_heading('美莲花美食双时段对比分析报告', 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
            
            # 定义字体设置函数
            def set_multilingual_font(run, size=12):
                """为run设置多语言字体支持"""
                # 设置主要字体（优先使用支持多语言的字体）
                try:
                    # 尝试使用Arial Unicode MS（如果系统有的话）
                    run.font.name = 'Arial Unicode MS'
                except:
                    # 备用方案：使用宋体
                    run.font.name = 'SimSun'
                
                run.font.size = Pt(size)
                
                # 设置不同语言的字体
                run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')  # 中文字体
                run.element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')  # 英文字体
                run.element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')  # 高位ANSI字符
                run.element.rPr.rFonts.set(qn('w:cs'), 'Arial Unicode MS')  # 复杂脚本（维语、阿拉伯文等）
            
            # 设置标题字体
            title_run = title_para.runs[0]
            set_multilingual_font(title_run, 18)
            
            # 添加时间段信息
            time_heading = doc.add_heading('分析时间段', level=1)
            time_heading_run = time_heading.runs[0]
            set_multilingual_font(time_heading_run, 14)
            
            # 时段信息段落
            p1_para = doc.add_paragraph(f'时段1：{p1_start} 至 {p1_end}')
            p1_run = p1_para.runs[0]
            set_multilingual_font(p1_run, 12)
            
            p2_para = doc.add_paragraph(f'时段2：{p2_start} 至 {p2_end}')
            p2_run = p2_para.runs[0]
            set_multilingual_font(p2_run, 12)
            
            # 添加分析内容标题
            content_heading = doc.add_heading('对比分析结果', level=1)
            content_heading_run = content_heading.runs[0]
            set_multilingual_font(content_heading_run, 14)
            
            # 将分析内容按行分割并添加到文档
            lines = analysis_content.split('\n')
            for line in lines:
                if line.strip():
                    para = doc.add_paragraph()
                    run = para.add_run(line.strip())
                    set_multilingual_font(run, 12)
            
            # 保存文档
            doc.save(file_path)
            
            InfoBar.success(
                title="导出成功",
                content=f"双时段对比报告已成功导出到：\n{file_path}",
                duration=3000,
                position=InfoBarPosition.TOP,
                parent=self
            )
            
        except ImportError:
            # 如果没有安装python-docx，使用纯文本文件作为备用方案
            txt_file_path = file_path.replace('.docx', '.txt')
            
            with open(txt_file_path, 'w', encoding='utf-8') as f:
                f.write('美莲花美食双时段对比分析报告\n')
                f.write('=' * 40 + '\n\n')
                f.write(f'时段1：{p1_start} 至 {p1_end}\n')
                f.write(f'时段2：{p2_start} 至 {p2_end}\n\n')
                f.write('对比分析结果：\n')
                f.write('-' * 20 + '\n')
                f.write(analysis_content)
            
            InfoBar.success(
                title="导出成功",
                content=f"双时段对比报告已成功导出到：\n{txt_file_path}\n\n注意：由于未安装python-docx库，已导出为文本文件格式。",
                duration=4000,
                position=InfoBarPosition.TOP,
                parent=self
            )

    def generate_pie_chart(self):
        """
        生成支出分类饼图
        """
        try:
            # 检查是否有数据且为单时段分析模式
            if not hasattr(self, 'current_expenses_data') or not self.current_expenses_data:
                InfoBar.warning(
                    title="图表提醒",
                    content="请先生成单时段分析报表",
                    duration=3000,
                    position=InfoBarPosition.TOP,
                    parent=self
                )
                return
            
            # 导入matplotlib
            try:
                import matplotlib.pyplot as plt
                import matplotlib
                
                # 设置中文字体
                matplotlib.rcParams['font.sans-serif'] = ['SimHei', 'DejaVu Sans']
                matplotlib.rcParams['axes.unicode_minus'] = False
                
            except ImportError:
                InfoBar.error(
                    title="功能不可用",
                    content="需要安装 matplotlib 库才能生成图表\n请运行: pip install matplotlib",
                    duration=5000,
                    position=InfoBarPosition.TOP,
                    parent=self
                )
                return
            
            # 按分类统计数据
            category_stats = {}
            for expense in self.current_expenses_data:
                category = expense['category_name']
                if category not in category_stats:
                    category_stats[category] = 0
                category_stats[category] += expense['amount']
            
            # 检查是否有有效数据
            if not category_stats:
                InfoBar.warning(
                    title="图表提醒",
                    content="没有找到有效的分类数据",
                    duration=3000,
                    position=InfoBarPosition.TOP,
                    parent=self
                )
                return
            
            # 准备图表数据
            categories = list(category_stats.keys())
            amounts = list(category_stats.values())
            
            # 创建饼图
            plt.figure(figsize=(10, 8))
            
            # 生成颜色
            colors = plt.cm.Set3(range(len(categories)))
            
            # 绘制饼图
            wedges, texts, autotexts = plt.pie(
                amounts, 
                labels=categories,
                autopct='%1.1f%%',
                startangle=90,
                colors=colors,
                textprops={'fontsize': 10}
            )
            
            # 设置标题
            plt.title('美莲花美食支出分类统计图', fontsize=16, fontweight='bold', pad=20)
            
            # 确保饼图是圆形
            plt.axis('equal')
            
            # 添加图例
            plt.legend(wedges, [f'{cat}: ¥{amt:.2f}' for cat, amt in zip(categories, amounts)],
                      title="分类详情",
                      loc="center left",
                      bbox_to_anchor=(1, 0, 0.5, 1))
            
            # 调整布局
            plt.tight_layout()
            
            # 显示图表
            plt.show()
            
            InfoBar.success(
                title="图表生成",
                content="饼图已成功生成并显示",
                duration=2000,
                position=InfoBarPosition.TOP,
                parent=self
            )
            
        except Exception as e:
            InfoBar.error(
                title="图表生成失败",
                content=f"生成饼图时发生错误：{str(e)}",
                duration=4000,
                position=InfoBarPosition.TOP,
                parent=self
            )

    def get_category_ug_name(self, category_cn_name):
        """
        根据中文分类名称获取维语分类名称
        """
        for cat_id, cat_info in self.categories.items():
            if cat_info['name_cn'] == category_cn_name:
                return cat_info['name_ug']
        return "نامەلۇم تۈر" 